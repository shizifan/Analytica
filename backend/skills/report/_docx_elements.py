"""DOCX element builder functions (python-docx).

Each function receives a ``Document`` and appends elements to it.
All are stateless — the orchestrator in ``docx_gen.py`` calls them in
the desired order.
"""
from __future__ import annotations

from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from backend.skills.report import _theme as T

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _rgb(t: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*t)


def _set_cell_shading(cell, hex_color: str):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _fmt_number(v: Any) -> str:
    """Format a numeric value for display."""
    if v is None:
        return "-"
    if isinstance(v, float):
        if abs(v) < 1:
            return f"{v:.4f}"
        return f"{v:,.2f}"
    return str(v)

# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------

def build_styles(doc: Document) -> None:
    """Configure document-wide styles for headings, body text and tables."""
    style = doc.styles["Normal"]
    style.font.name = T.FONT_CN
    style.font.size = Pt(T.SIZE_BODY)
    style.font.color.rgb = _rgb(T.RGB_NEUTRAL)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    # Set East Asia font
    style.element.rPr.rFonts.set(qn("w:eastAsia"), T.FONT_CN)

    for level, size, color in [
        ("Heading 1", T.SIZE_H1, T.RGB_PRIMARY),
        ("Heading 2", T.SIZE_H2, T.RGB_SECONDARY),
    ]:
        hs = doc.styles[level]
        hs.font.name = T.FONT_CN
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = _rgb(color)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), T.FONT_CN)

# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def build_cover_page(doc: Document, title: str, author: str, date: str) -> None:
    """Add a styled cover page followed by a page break."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = Pt(T.SIZE_TITLE)
    run.font.bold = True
    run.font.color.rgb = _rgb(T.RGB_PRIMARY)
    run.font.name = T.FONT_CN

    # Accent line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("━" * 30)
    run_line.font.color.rgb = _rgb(T.RGB_ACCENT)
    run_line.font.size = Pt(14)

    # Author / date
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"编制：{author}    日期：{date}")
    run_meta.font.size = Pt(T.SIZE_BODY)
    run_meta.font.color.rgb = _rgb(T.RGB_NEUTRAL)

    doc.add_page_break()

# ---------------------------------------------------------------------------
# TOC placeholder
# ---------------------------------------------------------------------------

def build_toc_placeholder(doc: Document) -> None:
    """Insert a TOC field code that Word can refresh on open."""
    heading = doc.add_heading("目  录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)
    run2 = p.add_run(' TOC \\o "1-2" \\h \\z \\u ')
    run2._element.tag = qn("w:instrText")
    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_char_end)

    p_hint = doc.add_paragraph('（请在 Word 中右键此处，选择"更新域"以生成目录）')
    p_hint.runs[0].font.size = Pt(T.SIZE_SMALL)
    p_hint.runs[0].font.color.rgb = _rgb(T.RGB_NEUTRAL)
    p_hint.runs[0].italic = True

    doc.add_page_break()

# ---------------------------------------------------------------------------
# Section heading
# ---------------------------------------------------------------------------

def build_section_heading(doc: Document, number: int, title: str) -> None:
    """Add a section heading. ``number`` is kept for backward compatibility
    with pre-batch-4 agent calls but no longer rendered — Chinese section
    names already carry "一、/二、" prefixes; stacking "1." on top produced
    visible double-numbering ("1. 一、经营摘要") in the pre-batch-4 output.
    """
    _ = number
    doc.add_heading(title, level=1)


def build_kpi_row(doc: Document, kpis: list) -> None:
    """Render a KPI card row as a single-row docx table.

    Each cell shows ``label`` / ``value`` / ``sub`` stacked vertically, with
    the accent-coloured value font weighted. No-op on empty list.
    """
    if not kpis:
        return
    tbl = doc.add_table(rows=1, cols=len(kpis))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, k in enumerate(kpis):
        cell = tbl.cell(0, i)
        # Cell shading
        _set_cell_shading(cell, T.BG_LIGHT)

        # Label
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.add_run(k.label)
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = _rgb(T.RGB_NEUTRAL)

        # Value
        p_val = cell.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p_val.add_run(k.value)
        run_val.bold = True
        run_val.font.size = Pt(18)
        if k.trend == "positive":
            run_val.font.color.rgb = _rgb(T.RGB_POSITIVE)
        elif k.trend == "negative":
            run_val.font.color.rgb = _rgb(T.RGB_NEGATIVE)
        else:
            run_val.font.color.rgb = _rgb(T.RGB_PRIMARY)

        # Sub
        if k.sub:
            p_sub = cell.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(k.sub)
            run_sub.font.size = Pt(8)
            run_sub.font.color.rgb = _rgb(T.RGB_NEUTRAL)

    # Space after the KPI row
    doc.add_paragraph("")

# ---------------------------------------------------------------------------
# Narrative paragraphs
# ---------------------------------------------------------------------------

def build_narrative(doc: Document, text: str) -> None:
    """Add narrative paragraphs, splitting on double-newlines."""
    for para_text in text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        run = p.add_run(para_text)
        run.font.size = Pt(T.SIZE_BODY)
        run.font.name = T.FONT_CN

# ---------------------------------------------------------------------------
# Stats table
# ---------------------------------------------------------------------------

def build_stats_table(doc: Document, summary_stats: dict[str, Any]) -> None:
    """Render summary_stats dict as a formatted Word table."""
    if not summary_stats:
        return

    doc.add_heading("统计数据概览", level=2)

    # Detect structure: {col: {metric: val}} or {col: {sub_col: {metric: val}}}
    first_val = next(iter(summary_stats.values()))
    if isinstance(first_val, dict) and not any(k in first_val for k in ("mean", "median", "std", "min", "max")):
        # Nested: grouped stats — flatten
        flat: dict[str, dict] = {}
        for group_key, cols in summary_stats.items():
            if isinstance(cols, dict):
                for col_name, metrics in cols.items():
                    flat[f"{group_key}/{col_name}"] = metrics if isinstance(metrics, dict) else {}
        summary_stats = flat if flat else summary_stats

    metrics = ["mean", "median", "std", "min", "max"]
    headers = ["指标"] + [m for m in metrics]
    rows_data: list[list[str]] = []
    for col, vals in summary_stats.items():
        if not isinstance(vals, dict):
            continue
        row = [str(col)]
        for m in metrics:
            row.append(_fmt_number(vals.get(m)))
        rows_data.append(row)

    if not rows_data:
        return

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, T.PRIMARY)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(T.SIZE_TABLE_HEADER)
                run.font.color.rgb = _rgb(T.RGB_WHITE)
                run.font.name = T.FONT_NUM

    # Data rows
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            if ri % 2 == 1:
                _set_cell_shading(cell, T.BG_LIGHT)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(T.SIZE_TABLE_BODY)
                    run.font.name = T.FONT_NUM if ci > 0 else T.FONT_CN

    doc.add_paragraph("")

# ---------------------------------------------------------------------------
# Growth indicators
# ---------------------------------------------------------------------------

def build_growth_indicators(doc: Document, growth_rates: dict[str, dict[str, float | None]]) -> None:
    """Render growth rates as colored inline indicators."""
    if not growth_rates:
        return

    doc.add_heading("增长率指标", level=2)

    for col, rates in growth_rates.items():
        if not isinstance(rates, dict):
            continue
        yoy = rates.get("yoy")
        mom = rates.get("mom")

        parts: list[str] = []
        if yoy is not None:
            arrow = "\u2191" if yoy >= 0 else "\u2193"
            parts.append(f"同比 {arrow}{abs(yoy)*100:.2f}%")
        if mom is not None:
            arrow = "\u2191" if mom >= 0 else "\u2193"
            parts.append(f"环比 {arrow}{abs(mom)*100:.2f}%")

        if not parts:
            continue

        p = doc.add_paragraph()
        run_label = p.add_run(f"{col}：")
        run_label.font.bold = True
        run_label.font.size = Pt(T.SIZE_BODY)

        for part in parts:
            run_val = p.add_run(f"  {part}")
            run_val.font.size = Pt(T.SIZE_BODY)
            run_val.font.name = T.FONT_NUM
            if "\u2191" in part:
                run_val.font.color.rgb = _rgb(T.RGB_POSITIVE)
            else:
                run_val.font.color.rgb = _rgb(T.RGB_NEGATIVE)

    doc.add_paragraph("")

# ---------------------------------------------------------------------------
# DataFrame → table
# ---------------------------------------------------------------------------

def build_dataframe_table(doc: Document, df: pd.DataFrame, max_rows: int = 20) -> None:
    """Render a DataFrame as a Word table (capped at *max_rows*)."""
    if df is None or df.empty:
        return

    doc.add_heading("数据明细", level=2)

    display_df = df.head(max_rows)
    n_rows, n_cols = display_df.shape

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for ci, col_name in enumerate(display_df.columns):
        cell = table.rows[0].cells[ci]
        cell.text = str(col_name)
        _set_cell_shading(cell, T.PRIMARY)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(T.SIZE_TABLE_HEADER)
                run.font.color.rgb = _rgb(T.RGB_WHITE)

    # Data
    for ri in range(n_rows):
        for ci in range(n_cols):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = _fmt_number(display_df.iloc[ri, ci])
            if ri % 2 == 1:
                _set_cell_shading(cell, T.BG_LIGHT)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(T.SIZE_TABLE_BODY)
                    run.font.name = T.FONT_NUM

    if len(df) > max_rows:
        p = doc.add_paragraph(f"（仅展示前 {max_rows} 行，共 {len(df)} 行）")
        p.runs[0].font.size = Pt(T.SIZE_SMALL)
        p.runs[0].italic = True

    doc.add_paragraph("")

# ---------------------------------------------------------------------------
# Chart data table (ECharts option → Word table)
# ---------------------------------------------------------------------------

def build_chart_data_table(doc: Document, chart_option: dict[str, Any]) -> None:
    """Extract data from an ECharts option dict and render as a table."""
    x_axis = chart_option.get("xAxis", {})
    categories = x_axis.get("data", []) if isinstance(x_axis, dict) else []
    series_list = chart_option.get("series", [])

    if not categories or not series_list:
        return

    title_obj = chart_option.get("title", {})
    chart_title = title_obj.get("text", "图表数据") if isinstance(title_obj, dict) else "图表数据"
    doc.add_heading(chart_title, level=2)

    headers = ["类别"] + [s.get("name", f"系列{i+1}") for i, s in enumerate(series_list)]
    n_rows = len(categories)

    table = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = str(h)
        _set_cell_shading(cell, T.SECONDARY)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(T.SIZE_TABLE_HEADER)
                run.font.color.rgb = _rgb(T.RGB_WHITE)

    # Data
    for ri, cat in enumerate(categories):
        table.rows[ri + 1].cells[0].text = str(cat)
        for si, s in enumerate(series_list):
            data_list = s.get("data", [])
            val = data_list[ri] if ri < len(data_list) else None
            table.rows[ri + 1].cells[si + 1].text = _fmt_number(val)
        if ri % 2 == 1:
            for ci in range(len(headers)):
                _set_cell_shading(table.rows[ri + 1].cells[ci], T.BG_LIGHT)
        for ci in range(len(headers)):
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(T.SIZE_TABLE_BODY)
                    run.font.name = T.FONT_NUM

    doc.add_paragraph("")

# ---------------------------------------------------------------------------
# Page header & footer
# ---------------------------------------------------------------------------

def build_page_header_footer(doc: Document, title: str) -> None:
    """Add page header (report title) and footer (page number)."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = title
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.size = Pt(T.SIZE_SMALL)
        run.font.color.rgb = _rgb(T.RGB_NEUTRAL)

    # Footer — page number via field code
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_begin)
    run2 = fp.add_run(" PAGE ")
    run2._element.tag = qn("w:instrText")
    run3 = fp.add_run()
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_end)

# ---------------------------------------------------------------------------
# Summary section
# ---------------------------------------------------------------------------

def build_summary_section(doc: Document, items: list) -> None:
    """Build the final summary / recommendations section."""
    doc.add_heading("总结与建议", level=1)

    if not items:
        p = doc.add_paragraph("以上分析基于数据，仅供参考。")
        p.runs[0].font.color.rgb = _rgb(T.RGB_NEUTRAL)
        return

    for item in items:
        text = item.text if hasattr(item, "text") else str(item)
        p = doc.add_paragraph()
        run = p.add_run(f"  {text}")
        run.font.size = Pt(T.SIZE_BODY)
        p.paragraph_format.first_line_indent = Cm(0.75)
