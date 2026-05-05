"""DOCX element builder functions (python-docx).

Each function receives a ``Document`` and appends elements to it.
All are stateless — the orchestrator in ``docx_gen.py`` calls them in
the desired order.
"""
from __future__ import annotations

from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from backend.tools._field_labels import col_label, metric_label
from backend.tools.report import _theme as T
from backend.tools.report._theme import Theme, get_theme

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _rgb(t: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*t)


def _set_cell_shading(cell, hex_color: str):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _fmt_number(v: Any) -> str:
    """Format a numeric value for display."""
    if v is None:
        return "-"
    if isinstance(v, float):
        if abs(v) < 1:
            return f"{v:.4f}"
        return f"{v:,.2f}"
    return str(v)


def _hex_from_rgb(rgb: tuple[int, int, int]) -> str:
    """RGB tuple → 6-digit hex string (no '#')."""
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_small_caps(style_element) -> None:
    """Add <w:smallCaps/> to a style's <w:rPr>."""
    rPr = style_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_element.insert(0, rPr)
    caps = OxmlElement("w:smallCaps")
    caps.set(qn("w:val"), "true")
    rPr.append(caps)


def _set_paragraph_top_border(pPr_or_p, color: tuple[int, int, int],
                               sz: int = 8) -> None:
    """Set a 1pt top border on a paragraph."""
    if hasattr(pPr_or_p, "get_or_add_pPr"):
        pPr = pPr_or_p.get_or_add_pPr()
    elif pPr_or_p.tag == qn("w:pPr"):
        pPr = pPr_or_p
    else:
        pPr = pPr_or_p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPr_or_p.insert(0, pPr)
    hex_c = _hex_from_rgb(color)
    border_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="4" w:color="{hex_c}"/>'
        f'</w:pBdr>'
    )
    pPr.append(parse_xml(border_xml))


def _set_paragraph_bottom_border(pPr_or_p, color: tuple[int, int, int],
                                  sz: int = 8) -> None:
    """Set a 1pt bottom border on a paragraph."""
    if hasattr(pPr_or_p, "get_or_add_pPr"):
        pPr = pPr_or_p.get_or_add_pPr()
    elif pPr_or_p.tag == qn("w:pPr"):
        pPr = pPr_or_p
    else:
        pPr = pPr_or_p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPr_or_p.insert(0, pPr)
    hex_c = _hex_from_rgb(color)
    border_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{hex_c}"/>'
        f'</w:pBdr>'
    )
    pPr.append(parse_xml(border_xml))


def _set_paragraph_left_border(pPr_or_p, color: tuple[int, int, int],
                               sz: int = 18) -> None:
    """Set a left border on a paragraph."""
    if hasattr(pPr_or_p, "get_or_add_pPr"):
        pPr = pPr_or_p.get_or_add_pPr()
    elif pPr_or_p.tag == qn("w:pPr"):
        pPr = pPr_or_p
    else:
        pPr = pPr_or_p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPr_or_p.insert(0, pPr)
    hex_c = _hex_from_rgb(color)
    border_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="8" w:color="{hex_c}"/>'
        f'</w:pBdr>'
    )
    pPr.append(parse_xml(border_xml))


def _set_paragraph_shading(pPr_or_p, color: tuple[int, int, int]) -> None:
    """Set paragraph background shading."""
    if hasattr(pPr_or_p, "get_or_add_pPr"):
        pPr = pPr_or_p.get_or_add_pPr()
    elif pPr_or_p.tag == qn("w:pPr"):
        pPr = pPr_or_p
    else:
        pPr = pPr_or_p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPr_or_p.insert(0, pPr)
    hex_c = _hex_from_rgb(color)
    shading_xml = (
        f'<w:shd {nsdecls("w")} w:val="clear" '
        f'w:color="auto" w:fill="{hex_c}"/>'
    )
    pPr.append(parse_xml(shading_xml))

# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------

def build_styles(doc: Document, theme: Theme | None = None) -> None:
    """Configure document-wide styles per Liangang Data Journal spec (SS5.1).

    Style table:

    - Normal:  font_body 11pt #1F1A12  1.5x line spacing  2-char indent
    - Title:   font_display 32pt bold #004889  left  bottom 1pt navy border
    - Subtitle:font_display italic 18pt #5E5648  after 24pt
    - Heading 1: 22pt bold #1F1A12  before 24pt  top 1pt navy border
    - Heading 2: 16pt bold #1F1A12  before 16pt
    - Heading 3: 14pt bold #1F1A12  before 12pt
    - Kicker (char):  font_ui 9pt #004889  smcp
    - Lede:     font_display italic 14pt #1F1A12  no indent  after 12pt
    - Narrative:font_body 11pt #1F1A12  1.5x  2-char indent
    - Callout:  font_body 11pt #1F1A12  left 3pt bronze border + paper bg
    - Pullquote:font_display italic 14pt #1F1A12  left 4pt bronze border
    - KPIValue (char): font_mono 28pt #004889
    - KPILabel (char): font_ui 8pt #9A8E78
    - Source:   font_ui italic 9pt #9A8E78  before 6pt
    - TableHeader (char): font_ui 9pt #9A8E78
    - TableNum (char):     font_mono 10pt #1F1A12
    """
    th = theme or get_theme()

    def _set_fonts(style, font_name, size_pt, color_rgb,
                   bold=False, italic=False):
        style.font.name = font_name
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = _rgb(color_rgb)
        if style.element.rPr is not None:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # -- Built-in paragraph styles ----------------------------------

    normal = doc.styles["Normal"]
    _set_fonts(normal, th.font_display, th.size_body, th.text_dark)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.75)

    title_st = doc.styles["Title"]
    _set_fonts(title_st, th.font_display, th.size_title, th.primary, bold=True)
    title_st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_st.paragraph_format.space_after = Pt(8)
    _set_paragraph_bottom_border(title_st.element, th.primary)

    subtitle_st = doc.styles["Subtitle"]
    _set_fonts(subtitle_st, th.font_display, 18, th.secondary, italic=True)
    subtitle_st.paragraph_format.space_after = Pt(24)

    h1 = doc.styles["Heading 1"]
    _set_fonts(h1, th.font_display, th.size_h1, th.text_dark, bold=True)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    _set_paragraph_top_border(h1.element, th.primary)

    h2 = doc.styles["Heading 2"]
    _set_fonts(h2, th.font_display, th.size_h2, th.text_dark, bold=True)
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    _set_fonts(h3, th.font_display, th.size_h3, th.text_dark, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # -- Custom paragraph styles ------------------------------------

    def _add_or_get_para(name):
        try:
            return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            return doc.styles[name]

    lede = _add_or_get_para("Lede")
    _set_fonts(lede, th.font_display, 14, th.text_dark, italic=True)
    lede.paragraph_format.first_line_indent = Cm(0)
    lede.paragraph_format.space_after = Pt(12)
    lede.paragraph_format.space_before = Pt(4)

    narrative = _add_or_get_para("Narrative")
    _set_fonts(narrative, th.font_display, th.size_body, th.text_dark)
    narrative.paragraph_format.first_line_indent = Cm(0.75)
    narrative.paragraph_format.line_spacing = 1.5
    narrative.paragraph_format.space_after = Pt(4)

    callout = _add_or_get_para("Callout")
    _set_fonts(callout, th.font_display, th.size_body, th.text_dark)
    callout.paragraph_format.first_line_indent = Cm(0)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.space_before = Pt(4)
    _set_paragraph_left_border(callout.element, th.accent, 18)
    _set_paragraph_shading(callout.element, th.bg_light)

    pullquote = _add_or_get_para("Pullquote")
    _set_fonts(pullquote, th.font_display, 14, th.text_dark, italic=True)
    pullquote.paragraph_format.first_line_indent = Cm(0)
    pullquote.paragraph_format.left_indent = Cm(1.5)
    pullquote.paragraph_format.space_after = Pt(12)
    _set_paragraph_left_border(pullquote.element, th.accent, 24)

    source = _add_or_get_para("Source")
    _set_fonts(source, th.font_ui, th.size_small, th.neutral, italic=True)
    source.paragraph_format.first_line_indent = Cm(0)
    source.paragraph_format.space_before = Pt(6)
    source.paragraph_format.space_after = Pt(4)

    # -- Custom character styles ------------------------------------

    def _add_or_get_char(name):
        try:
            return doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        except ValueError:
            return doc.styles[name]

    kicker = _add_or_get_char("Kicker")
    _set_fonts(kicker, th.font_ui, th.size_small, th.primary)
    _set_small_caps(kicker.element)

    kpi_value_st = _add_or_get_char("KPIValue")
    _set_fonts(kpi_value_st, th.font_num, th.size_kpi_large, th.primary)

    kpi_label_st = _add_or_get_char("KPILabel")
    _set_fonts(kpi_label_st, th.font_ui, th.size_kpi_label, th.neutral)
    _set_small_caps(kpi_label_st.element)

    th_header = _add_or_get_char("TableHeader")
    _set_fonts(th_header, th.font_ui, th.size_table_header, th.neutral)

    th_num = _add_or_get_char("TableNum")
    _set_fonts(th_num, th.font_num, th.size_table_body, th.text_dark)

# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def build_cover_page(doc: Document, title: str, author: str, date: str,
                     *, theme: Theme | None = None) -> None:
    """Liangang Data Journal cover page (SS5.6).

    Layout: paper background, top bronze brand bar, navy title 32pt,
    deck italic 18pt ink_2, bronze endmark, metadata small text.
    """
    th = theme or get_theme()

    # Top spacer
    for _ in range(3):
        doc.add_paragraph("")

    # Bronze brand bar (top accent stripe)
    p_bar = doc.add_paragraph()
    p_bar.paragraph_format.space_after = Pt(0)
    pPr = p_bar._p.get_or_add_pPr()
    _set_paragraph_shading(p_bar._p, th.accent)
    # Make the bar paragraph a fixed small height via spacing
    bar_run = p_bar.add_run(" ")
    bar_run.font.size = Pt(8)
    bar_run.font.color.rgb = _rgb(th.accent)

    # Spacer
    doc.add_paragraph("")

    # Title — 32pt navy display font
    p_title = doc.add_paragraph()
    p_title.style = doc.styles["Title"]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(th.size_title)
    run_title.font.bold = True
    run_title.font.color.rgb = _rgb(th.primary)
    run_title.font.name = th.font_display
    if p_title.paragraph_format.first_line_indent:
        p_title.paragraph_format.first_line_indent = Cm(0)

    # Deck / subtitle
    if author:
        p_deck = doc.add_paragraph()
        p_deck.style = doc.styles["Subtitle"]
        p_deck.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_deck = p_deck.add_run(f"{author}")
        run_deck.font.size = Pt(18)
        run_deck.font.italic = True
        run_deck.font.color.rgb = _rgb(th.secondary)
        run_deck.font.name = th.font_display

    # Spacer
    for _ in range(2):
        doc.add_paragraph("")

    # Metadata: date in kicker style
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_meta = p_meta.add_run(f"日期：{date}")
    run_meta.font.size = Pt(th.size_small)
    run_meta.font.color.rgb = _rgb(th.accent)
    run_meta.font.name = th.font_ui
    _set_small_caps_on_run(run_meta)

    # Endmark — bronze square (U+25A0 ■)
    p_endmark = doc.add_paragraph()
    p_endmark.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_endmark = p_endmark.add_run("\u25A0")
    run_endmark.font.size = Pt(10)
    run_endmark.font.color.rgb = _rgb(th.accent)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# TOC placeholder
# ---------------------------------------------------------------------------

def build_toc_placeholder(doc: Document, *, theme: Theme | None = None) -> None:
    """Insert a TOC field code that Word can refresh on open."""
    th = theme or get_theme()

    heading = doc.add_heading("目  录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)
    run2 = p.add_run(' TOC \\o "1-2" \\h \\z \\u ')
    run2._element.tag = qn("w:instrText")
    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_char_end)

    p_hint = doc.add_paragraph('（请在 Word 中右键此处，选择"更新域"以生成目录）')
    p_hint.runs[0].font.size = Pt(th.size_small)
    p_hint.runs[0].font.color.rgb = _rgb(th.neutral)
    p_hint.runs[0].italic = True

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Section heading
# ---------------------------------------------------------------------------

def build_section_heading(doc: Document, number: int, title: str,
                          *, theme: Theme | None = None) -> None:
    """Liangang Data Journal section divider (SS5.6).

    Uses Heading 1 style (22pt navy, top hairline via style).
    ``number`` renders as a roman numeral prefix when > 0.
    """
    th = theme or get_theme()

    # Roman numeral prefix (e.g. "I. ", "II. ") in bronze
    if number > 0:
        roman = _to_roman(number)
        p_roman = doc.add_paragraph()
        p_roman.paragraph_format.first_line_indent = Cm(0)
        p_roman.paragraph_format.space_after = Pt(0)
        p_roman.paragraph_format.space_before = Pt(16)
        run_roman = p_roman.add_run(f"{roman}")
        run_roman.font.size = Pt(36)
        run_roman.font.italic = True
        run_roman.font.color.rgb = _rgb(th.accent)
        run_roman.font.name = th.font_display

    # Chapter title as H1
    doc.add_heading(title, level=1)


def _to_roman(n: int) -> str:
    """1.. → I, II, III, IV, V..."""
    vals = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    result = ""
    for v, r in vals:
        while n >= v:
            result += r
            n -= v
    return result


def _set_small_caps_on_run(run) -> None:
    """Add <w:smallCaps/> to a single run's <w:rPr>."""
    rPr = run._element.get_or_add_rPr()
    caps = OxmlElement("w:smallCaps")
    caps.set(qn("w:val"), "true")
    rPr.append(caps)


# ---------------------------------------------------------------------------
# KPI row (legacy: multi-card overview row; kept for backward compat)
# ---------------------------------------------------------------------------

def build_kpi_row(doc: Document, kpis: list) -> None:
    """Render a KPI card row as a single-row docx table.

    Each cell shows ``label`` / ``value`` / ``sub`` stacked vertically, with
    the accent-coloured value font weighted. No-op on empty list.
    """
    if not kpis:
        return
    th = get_theme()
    tbl = doc.add_table(rows=1, cols=len(kpis))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, k in enumerate(kpis):
        cell = tbl.cell(0, i)
        _set_cell_shading(cell, th.hex_bg_light)

        # Label
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p_label.add_run(k.label)
        run_label.font.size = Pt(th.size_kpi_label)
        run_label.font.color.rgb = _rgb(th.neutral)
        run_label.font.name = th.font_ui

        # Value
        p_val = cell.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p_val.add_run(k.value)
        run_val.bold = True
        run_val.font.size = Pt(th.size_kpi_large)
        run_val.font.name = th.font_num
        if k.trend == "positive":
            run_val.font.color.rgb = _rgb(th.positive)
        elif k.trend == "negative":
            run_val.font.color.rgb = _rgb(th.negative)
        else:
            run_val.font.color.rgb = _rgb(th.primary)

        # Sub
        if k.sub:
            p_sub = cell.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(k.sub)
            run_sub.font.size = Pt(th.size_small)
            run_sub.font.color.rgb = _rgb(th.neutral)
            run_sub.font.name = th.font_ui

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# KPI strip (4-cell trend strip — Liangang Data Journal SS5.3)
# ---------------------------------------------------------------------------

def build_kpi_strip(
    doc: Document,
    items: tuple | list,
    *,
    theme: Theme | None = None,
) -> None:
    """Render a 4-cell KPI strip (起点 / 高点 / 当前 / 变化).

    Layout: 1x4 table, no cell borders. Table top/bottom 1pt navy
    paragraph borders. Each cell: label (smcp bronze), value (mono
    28pt navy/accent_dark), sub (ui 9pt neutral).
    """
    if not items or len(items) != 4:
        return

    th = theme or get_theme()
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove table style border
    tbl.style = "Table Grid"
    _remove_table_borders(tbl)

    for i, k in enumerate(items):
        cell = tbl.cell(0, i)
        cell.paragraphs[0].clear()

        # Label — smcp bronze, 9pt
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_label.paragraph_format.space_after = Pt(2)
        run_label = p_label.add_run(str(getattr(k, "label", "")))
        run_label.font.size = Pt(th.size_kpi_label)
        run_label.font.color.rgb = _rgb(th.accent)
        run_label.font.name = th.font_ui
        _set_small_caps_on_run(run_label)

        # Value — mono 28pt navy/gain or accent_dark/loss
        p_val = cell.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_val.paragraph_format.space_after = Pt(0)
        run_val = p_val.add_run(str(getattr(k, "value", "")))
        run_val.font.size = Pt(th.size_kpi_large)
        run_val.font.name = th.font_num
        trend = getattr(k, "trend", "")
        if trend == "gain":
            run_val.font.color.rgb = _rgb(th.positive)  # navy = primary
        elif trend == "loss":
            run_val.font.color.rgb = _rgb(th.negative)   # accent_dark
        else:
            run_val.font.color.rgb = _rgb(th.primary)

        # Sub — ui 9pt neutral
        sub_val = str(getattr(k, "sub", ""))
        if sub_val:
            p_sub = cell.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = Pt(2)
            run_sub = p_sub.add_run(sub_val)
            run_sub.font.size = Pt(th.size_small)
            run_sub.font.color.rgb = _rgb(th.neutral)
            run_sub.font.name = th.font_ui

    # Add top/bottom navy hairline borders to the paragraph
    _set_paragraph_top_border(doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph("")._element,
                              th.primary, sz=8)
    doc.add_paragraph("")


def _remove_table_borders(tbl) -> None:
    """Remove all cell borders from a table."""
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "nil")
                borders.append(e)
            tcPr.append(borders)


# ---------------------------------------------------------------------------
# Narrative paragraphs
# ---------------------------------------------------------------------------

def build_narrative(doc: Document, text: str,
                    *, theme: Theme | None = None) -> None:
    """Add narrative paragraphs using the 'Narrative' style (SS5.1)."""
    th = theme or get_theme()
    for para_text in text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(style="Narrative")
        # python-docx add_paragraph(style=...) may not apply run-level
        # formatting; explicitly set it.
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(para_text)
        run.font.size = Pt(th.size_body)
        run.font.name = th.font_display
        run.font.color.rgb = _rgb(th.text_dark)


def build_lede(doc: Document, text: str,
               *, theme: Theme | None = None) -> None:
    """Add a lede paragraph with drop cap (SS5.4)."""
    th = theme or get_theme()
    p = doc.add_paragraph(style="Lede")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.name = th.font_display
    run.font.color.rgb = _rgb(th.text_dark)
    _apply_dropcap(p, th)


def _apply_dropcap(para, theme: Theme, lines: int = 3) -> None:
    """Apply drop-cap to the first character of a paragraph via OOXML.

    Adds <w:framePr w:dropCap="drop" w:lines="{lines}"/> to the
    paragraph properties. Increases the first run's font size.
    LibreOffice renders framePr differently — graceful degradation
    to plain italic.
    """
    pPr = para._p.get_or_add_pPr()
    framePr = OxmlElement("w:framePr")
    framePr.set(qn("w:dropCap"), "drop")
    framePr.set(qn("w:lines"), str(lines))
    framePr.set(qn("w:wrap"), "around")
    framePr.set(qn("w:vAnchor"), "text")
    framePr.set(qn("w:hAnchor"), "text")
    # Remove existing framePr if any
    existing = pPr.find(qn("w:framePr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(framePr)
    if para.runs:
        para.runs[0].font.size = Pt(36)
        para.runs[0].font.color.rgb = _rgb(theme.primary)


# ---------------------------------------------------------------------------
# Callout
# ---------------------------------------------------------------------------

def build_callout(
    doc: Document, text: str, level: str = "warn", theme=None,
) -> None:
    """Phase 4.1 — render a callout block with left coloured border and
    light tinted background.

    ``level`` is one of ``"warn"`` / ``"info"``; the leading emoji and
    border colour pick from the theme's ``callout_*`` fields. The
    paragraph is added without first-line indent (callouts are
    standalone visual blocks, not inline narrative).
    """
    th = theme or get_theme()
    if level == "warn":
        emoji = "\u26A0"  # ⚠
        title_label = "\u6CE8\u610F"  # 注意
        bg = th.callout_warn_bg
        border = th.callout_warn_border
    else:  # info / fallback
        emoji = "\U0001F4A1"  # 💡
        title_label = "\u63D0\u793A"  # 提示
        bg = th.callout_info_bg
        border = th.callout_info_border

    border_hex = _hex_from_rgb(border)
    bg_hex = _hex_from_rgb(bg)

    p = doc.add_paragraph(style="Callout")
    p.paragraph_format.first_line_indent = Cm(0)
    pPr = p._p.get_or_add_pPr()
    # Left border + light fill — custom per level, overriding style default
    border_xml = f'''<w:pBdr {nsdecls("w")}>
        <w:left w:val="single" w:sz="24" w:space="6" w:color="{border_hex}"/>
    </w:pBdr>'''
    pPr.append(parse_xml(border_xml))
    shading_xml = (
        f'<w:shd {nsdecls("w")} w:val="clear" '
        f'w:color="auto" w:fill="{bg_hex}"/>'
    )
    pPr.append(parse_xml(shading_xml))

    # Title run (bold, coloured)
    title_run = p.add_run(f"{emoji} {title_label}\uFF1A")  # ：
    title_run.font.bold = True
    title_run.font.size = Pt(th.size_body)
    title_run.font.color.rgb = _rgb(border)
    title_run.font.name = th.font_display
    # Body run
    body_run = p.add_run(text)
    body_run.font.size = Pt(th.size_body)
    body_run.font.name = th.font_display
    body_run.font.color.rgb = _rgb(th.text_dark)

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Stats table
# ---------------------------------------------------------------------------

def build_stats_table(doc: Document, summary_stats: dict[str, Any]) -> None:
    """Render summary_stats dict as a formatted Word table."""
    if not summary_stats:
        return

    th = get_theme()
    doc.add_heading("\u7EDF\u8BA1\u6570\u636E\u6982\u89C8", level=2)  # 统计数据概览

    # Detect structure: {col: {metric: val}} or {col: {sub_col: {metric: val}}}
    first_val = next(iter(summary_stats.values()))
    if isinstance(first_val, dict) and not any(k in first_val for k in ("mean", "median", "std", "min", "max")):
        # Nested: grouped stats — flatten
        flat: dict[str, dict] = {}
        for group_key, cols in summary_stats.items():
            if isinstance(cols, dict):
                for col_name, metrics in cols.items():
                    flat[f"{group_key}/{col_name}"] = metrics if isinstance(metrics, dict) else {}
        summary_stats = flat if flat else summary_stats

    metrics = ["mean", "median", "std", "min", "max"]
    headers = ["\u6307\u6807"] + [metric_label(m) for m in metrics]  # 指标
    rows_data: list[list[str]] = []
    for col, vals in summary_stats.items():
        if not isinstance(vals, dict):
            continue
        row = [str(col)]
        for m in metrics:
            row.append(_fmt_number(vals.get(m)))
        rows_data.append(row)

    if not rows_data:
        return

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, th.hex_primary)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(th.size_table_header)
                run.font.color.rgb = _rgb(th.white)
                run.font.name = th.font_num

    # Data rows
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            if ri % 2 == 1:
                _set_cell_shading(cell, th.hex_bg_light)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(th.size_table_body)
                    run.font.name = th.font_num if ci > 0 else th.font_display

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Growth indicators
# ---------------------------------------------------------------------------

def build_growth_indicators(doc: Document, growth_rates: dict[str, dict[str, float | None]]) -> None:
    """Render growth rates as colored inline indicators."""
    if not growth_rates:
        return
    th = get_theme()

    doc.add_heading("\u589E\u957F\u7387\u6307\u6807", level=2)  # 增长率指标

    for col, rates in growth_rates.items():
        if not isinstance(rates, dict):
            continue
        yoy = rates.get("yoy")
        mom = rates.get("mom")

        parts: list[str] = []
        if yoy is not None:
            arrow = "\u2191" if yoy >= 0 else "\u2193"
            parts.append(f"\u540C\u6BD4 {arrow}{abs(yoy)*100:.2f}%")  # 同比
        if mom is not None:
            arrow = "\u2191" if mom >= 0 else "\u2193"
            parts.append(f"\u73AF\u6BD4 {arrow}{abs(mom)*100:.2f}%")  # 环比

        if not parts:
            continue

        p = doc.add_paragraph()
        run_label = p.add_run(f"{col_label(col)}：")
        run_label.font.bold = True
        run_label.font.size = Pt(th.size_body)

        for part in parts:
            run_val = p.add_run(f"  {part}")
            run_val.font.size = Pt(th.size_body)
            run_val.font.name = th.font_num
            if "\u2191" in part:
                run_val.font.color.rgb = _rgb(th.positive)
            else:
                run_val.font.color.rgb = _rgb(th.negative)

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# DataFrame → table
# ---------------------------------------------------------------------------

def build_dataframe_table(
    doc: Document,
    df: pd.DataFrame,
    max_rows: int = 20,
    *,
    highlight_rules: list[dict[str, Any]] | None = None,
    theme=None,
    heading: str = "\u6570\u636E\u660E\u7EC6",  # 数据明细
) -> None:
    """Render a DataFrame as a Word table (capped at *max_rows*).

    Phase 4.2 — ``highlight_rules`` paints body cells (row/col-based,
    see :mod:`backend.tools.report._table_highlight`) before the
    striped-row tint. When a rule applies, alternating-row shading is
    *replaced* by the rule colour for visual prominence.
    """
    if df is None or df.empty:
        return

    from backend.tools.report._table_highlight import (
        resolve_cell_highlights,
        rgb_to_hex,
    )

    th = theme or get_theme()
    doc.add_heading(heading, level=2)

    display_df = df.head(max_rows)
    n_rows, n_cols = display_df.shape

    headers = [str(c) for c in display_df.columns]
    cell_colors = resolve_cell_highlights(
        headers, n_rows, highlight_rules or [], th,
    )

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for ci, col_name in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = col_name
        _set_cell_shading(cell, th.hex_primary)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(th.size_table_header)
                run.font.color.rgb = _rgb(th.white)

    # Data
    for ri in range(n_rows):
        for ci in range(n_cols):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = _fmt_number(display_df.iloc[ri, ci])
            highlight = cell_colors.get((ri, ci))
            if highlight is not None:
                _set_cell_shading(cell, "#" + rgb_to_hex(highlight))
            elif ri % 2 == 1:
                _set_cell_shading(cell, th.hex_bg_light)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(th.size_table_body)
                    run.font.name = th.font_num
                    if highlight is not None:
                        lum = (
                            0.299 * highlight[0]
                            + 0.587 * highlight[1]
                            + 0.114 * highlight[2]
                        )
                        run.font.color.rgb = (
                            _rgb(th.white) if lum < 140
                            else _rgb(th.text_dark)
                        )
                        run.font.bold = True

    if len(df) > max_rows:
        p = doc.add_paragraph(
            f"\uFF08\u4EC5\u5C55\u793A\u524D {max_rows} \u884C\uFF0C"
            f"\u5171 {len(df)} \u884C\uFF09"
        )
        p.runs[0].font.size = Pt(th.size_small)
        p.runs[0].italic = True

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Chart data table (ECharts option -> Word table)
# ---------------------------------------------------------------------------

def build_chart_data_table(doc: Document, chart_option: dict[str, Any]) -> None:
    """Extract data from an ECharts option dict and render as a table."""
    x_axis = chart_option.get("xAxis", {})
    categories = x_axis.get("data", []) if isinstance(x_axis, dict) else []
    series_list = chart_option.get("series", [])

    if not categories or not series_list:
        return

    title_obj = chart_option.get("title", {})
    chart_title = title_obj.get("text", "\u56FE\u8868\u6570\u636E") if isinstance(title_obj, dict) else "\u56FE\u8868\u6570\u636E"
    doc.add_heading(chart_title, level=2)

    headers = ["\u7C7B\u522B"] + [s.get("name", f"\u7CFB\u5217{i+1}") for i, s in enumerate(series_list)]
    n_rows = len(categories)

    table = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = str(h)
        _set_cell_shading(cell, T.SECONDARY)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(T.SIZE_TABLE_HEADER)
                run.font.color.rgb = _rgb(T.RGB_WHITE)

    # Data
    for ri, cat in enumerate(categories):
        table.rows[ri + 1].cells[0].text = str(cat)
        for si, s in enumerate(series_list):
            data_list = s.get("data", [])
            val = data_list[ri] if ri < len(data_list) else None
            table.rows[ri + 1].cells[si + 1].text = _fmt_number(val)
        if ri % 2 == 1:
            for ci in range(len(headers)):
                _set_cell_shading(table.rows[ri + 1].cells[ci], T.BG_LIGHT)
        for ci in range(len(headers)):
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(T.SIZE_TABLE_BODY)
                    run.font.name = T.FONT_NUM

    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# Hairline table (Data Journal editorial style — SS5.2)
# ---------------------------------------------------------------------------

def build_hairline_table(
    doc: Document,
    df: pd.DataFrame,
    *,
    theme: Theme | None = None,
    heading: str = "",
    source: str = "",
    max_rows: int = 20,
    even_row_shading: bool = False,
) -> None:
    """Render a DataFrame as a hairline table per Liangang Data Journal spec.

    - No Table Grid style — only top/bottom hairline on header and footer rows.
    - Header: top 1pt text_dark, bottom 0.5pt rule_soft.
    - Last data row: bottom 1pt text_dark.
    - Number columns: right-aligned, mono font.
    - Optional even-row paper_2 shading for long tables.
    """
    if df is None or df.empty:
        return

    th = theme or get_theme()
    display_df = df.head(max_rows)
    n_rows, n_cols = display_df.shape
    headers = [str(c) for c in display_df.columns]

    if heading:
        doc.add_heading(heading, level=3)

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove "Table Grid" border
    _remove_table_borders(table)

    # Determine numeric columns
    num_cols = set()
    for ci, col in enumerate(display_df.columns):
        if pd.api.types.is_numeric_dtype(display_df[col]):
            num_cols.add(ci)

    # Header row — top border, bottom hairline
    for ci, col_name in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = col_name
        # Top border (1pt text_dark)
        _set_cell_top_border(cell, th.text_dark, sz=8)
        # Bottom border (hairline rule_soft ~0.5pt)
        _set_cell_bottom_border(cell, th.text_dark, sz=4,
                                opacity=20)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci in num_cols else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(th.size_table_header)
                run.font.name = th.font_ui
                run.font.color.rgb = _rgb(th.neutral)
                _set_small_caps_on_run(run)

    # Data rows
    for ri in range(n_rows):
        is_last = (ri == n_rows - 1)
        for ci in range(n_cols):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = _fmt_number(display_df.iloc[ri, ci])
            # Even row shading (optional)
            if even_row_shading and ri % 2 == 1:
                _set_cell_shading(cell, th.hex_bg_light)
            # Last row bottom border
            if is_last:
                _set_cell_bottom_border(cell, th.text_dark, sz=8)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci in num_cols else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(th.size_table_body)
                    run.font.name = th.font_num if ci in num_cols else th.font_display

    if source:
        p_src = doc.add_paragraph(style="Source")
        p_src.add_run(source)
        p_src.runs[-1].font.size = Pt(th.size_small)
        p_src.runs[-1].font.italic = True
        p_src.runs[-1].font.color.rgb = _rgb(th.neutral)

    if len(df) > max_rows:
        p = doc.add_paragraph(
            f"\uFF08\u4EC5\u5C55\u793A\u524D {max_rows} \u884C\uFF0C"
            f"\u5171 {len(df)} \u884C\uFF09"
        )
        p.runs[0].font.size = Pt(th.size_small)
        p.runs[0].italic = True

    doc.add_paragraph("")


def _set_cell_top_border(cell, color: tuple[int, int, int],
                         sz: int = 8, opacity: int = 100) -> None:
    """Set top border on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.insert(0, borders)
    hex_c = _hex_from_rgb(color)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(sz))
    top.set(qn("w:color"), hex_c)
    # Remove existing
    existing = borders.find(qn("w:top"))
    if existing is not None:
        borders.remove(existing)
    borders.append(top)


def _set_cell_bottom_border(cell, color: tuple[int, int, int],
                            sz: int = 8, opacity: int = 100) -> None:
    """Set bottom border on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.insert(0, borders)
    hex_c = _hex_from_rgb(color)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:color"), hex_c)
    existing = borders.find(qn("w:bottom"))
    if existing is not None:
        borders.remove(existing)
    borders.append(bottom)


# ---------------------------------------------------------------------------
# Hairline paragraph (divider)
# ---------------------------------------------------------------------------

def build_hairline_paragraph(doc: Document,
                             *, theme: Theme | None = None) -> None:
    """Insert a thin hairline divider paragraph (SS5.5)."""
    th = theme or get_theme()
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    _set_paragraph_top_border(p._p, th.text_dark, sz=2)  # hairline


# ---------------------------------------------------------------------------
# Appendix sub-heading
# ---------------------------------------------------------------------------

def build_appendix_subheading(
    doc: Document, label: str = "\u5B8C\u6574\u6570\u636E",
    *, theme: Theme | None = None,
) -> None:
    """Add a '完整数据' sub-heading for section appendix (SS5.5)."""
    th = theme or get_theme()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(label)
    run.font.size = Pt(th.size_h3)
    run.font.bold = True
    run.font.color.rgb = _rgb(th.secondary)
    run.font.name = th.font_ui

    # Add a bookmark anchor
    bookmark_id = f"_appendix_{label.replace(' ', '_')}"
    _add_bookmark(p, bookmark_id)


def _add_bookmark(paragraph, bookmark_name: str) -> None:
    """Insert an OOXML bookmark at the start of a paragraph."""
    p_elem = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "1")
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "1")
    # Insert at beginning of paragraph
    p_elem.insert(0, end)
    p_elem.insert(0, start)


# ---------------------------------------------------------------------------
# Page header & footer
# ---------------------------------------------------------------------------

def build_page_header_footer(doc: Document, title: str,
                             *, theme: Theme | None = None) -> None:
    """Add page header (kicker style) and footer (page number + endmark)."""
    th = theme or get_theme()
    section = doc.sections[0]

    # Header — left: kicker label, right: (empty or page field)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = hp.add_run(title)
    run_h.font.size = Pt(8)
    run_h.font.color.rgb = _rgb(th.accent)
    run_h.font.name = th.font_ui
    _set_small_caps_on_run(run_h)

    # Footer — page number centred
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = fp.add_run()
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_begin)
    run2 = fp.add_run(" PAGE ")
    run2._element.tag = qn("w:instrText")
    run3 = fp.add_run()
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_end)


# ---------------------------------------------------------------------------
# Summary section
# ---------------------------------------------------------------------------

def build_summary_section(doc: Document, items: list) -> None:
    """Build the final summary / recommendations section."""
    doc.add_heading("\u603B\u7ED3\u4E0E\u5EFA\u8BAE", level=1)  # 总结与建议

    if not items:
        p = doc.add_paragraph(
            "\u4EE5\u4E0A\u5206\u6790\u57FA\u4E8E\u6570\u636E\uFF0C"
            "\u4EC5\u4F9B\u53C2\u8003\u3002"
        )
        p.runs[0].font.color.rgb = _rgb(T.RGB_NEUTRAL)
        return

    for item in items:
        text = item.text if hasattr(item, "text") else str(item)
        p = doc.add_paragraph(style="Narrative")
        run = p.add_run(f"  {text}")
        run.font.size = Pt(T.SIZE_BODY)
