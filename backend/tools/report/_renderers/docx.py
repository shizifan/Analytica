"""DOCX BlockRenderer — Step 4.

Liangang Data Journal PR-3: full DOCX visual redesign with editorial
Data Journal styles, hairline tables, KPI strip, drop cap, and
section-level appendix buffering.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass

import pandas as pd

from backend.tools.report import _docx_elements as E
from backend.tools.report._block_renderer import BlockRendererBase
from backend.tools.report._outline import (
    Asset,
    ChartBlock,
    ChartTablePairBlock,
    ComparisonGridBlock,
    GrowthIndicatorsBlock,
    KpiRowBlock,
    KpiStripBlock,
    OutlineSection,
    ParagraphBlock,
    ReportOutline,
    SectionCoverBlock,
    StatsAsset,
    TableAsset,
    TableBlock,
)
from backend.tools.report._retry_config import RENDERER_OOM_RETRY, renderer_retry

logger = logging.getLogger("analytica.tools.report.docx")

# ── 图表渲染重试（matplotlib OOM） ──────────────────────────────

_retry_renderer_chart = renderer_retry(
    RENDERER_OOM_RETRY,
    on_exc=(MemoryError,),
    logger_name="analytica.tools.report.docx",
    fallback_value=None,
)


@dataclass
class _AppendixItem:
    """Duck-typed substitute for ``SummaryTextItem`` used by
    ``E.build_summary_section`` — the function only reads ``.text``."""

    text: str


@dataclass
class _AppendixTableEntry:
    """Buffered table + source reference for section appendix (SS5.5)."""
    block: TableBlock | ChartTablePairBlock
    asset: Asset


class DocxBlockRenderer(BlockRendererBase):
    _step_label = "Step 4"

    def __init__(self, theme=None) -> None:  # noqa: ANN001
        super().__init__(theme=theme)
        self._doc = None
        self._title: str = ""
        self._current_role: str = "status"
        self._current_section_name: str = ""
        # Paragraph appendix (legacy: for SummaryTextItem-style appendix)
        self._appendix_buffer: list = []
        # Section-level appendix: buffered table assets flushed per section
        self._section_appendix_buffer: dict[str, list[_AppendixTableEntry]] = {}

    # ---- Lifecycle -----------------------------------------------------

    def begin_document(self, outline: ReportOutline) -> None:
        from docx import Document
        self._doc = Document()
        self._title = outline.metadata.get("title", "")
        author = outline.metadata.get("author", "")
        date = outline.metadata.get("date", "")

        E.build_styles(self._doc, theme=self._theme)
        E.build_page_header_footer(self._doc, self._title, theme=self._theme)
        E.build_cover_page(self._doc, self._title, author, date, theme=self._theme)
        E.build_toc_placeholder(self._doc, theme=self._theme)
        if outline.kpi_summary:
            E.build_kpi_row(self._doc, list(outline.kpi_summary))

    def end_document(self) -> bytes:
        buffer = io.BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()

    def begin_section(self, section: OutlineSection, index: int) -> None:
        self._current_role = section.role
        self._current_section_name = section.name
        if section.role == "appendix":
            self._appendix_buffer = []
        # Initialize section appendix buffer for non-appendix sections
        if section.role != "appendix":
            self._section_appendix_buffer[section.name] = []

    def end_section(self, section: OutlineSection, index: int) -> None:
        # Legacy appendix path
        if section.role == "appendix":
            E.build_summary_section(self._doc, list(self._appendix_buffer))
            return

        # Flush section appendix: 完整数据 sub-section (SS5.5)
        buffered = self._section_appendix_buffer.get(section.name, [])
        if buffered:
            E.build_hairline_paragraph(self._doc, theme=self._theme)
            E.build_appendix_subheading(self._doc, theme=self._theme)
            for entry in buffered:
                if isinstance(entry.asset, TableAsset):
                    df = (
                        pd.DataFrame.from_records(entry.asset.df_records)
                        if entry.asset.df_records
                        else pd.DataFrame()
                    )
                    if not df.empty:
                        E.build_hairline_table(
                            self._doc, df, theme=self._theme,
                            heading=getattr(entry.block, "caption", ""),
                            source=getattr(entry.block, "source", ""),
                        )

    # ---- Block emitters ------------------------------------------------

    def emit_kpi_row(self, block: KpiRowBlock) -> None:
        if block.items:
            E.build_kpi_row(self._doc, list(block.items))

    def emit_kpi_strip(self, block: KpiStripBlock) -> None:
        """Render a 4-cell KPI strip (SS5.3)."""
        if block.items:
            E.build_kpi_strip(self._doc, block.items, theme=self._theme)

    def emit_paragraph(self, block: ParagraphBlock) -> None:
        if (
            self._current_role != "appendix"
            and block.style in ("callout-warn", "callout-info")
        ):
            level = "warn" if block.style == "callout-warn" else "info"
            E.build_callout(self._doc, block.text, level=level, theme=self._theme)
            return

        if self._current_role == "appendix":
            self._appendix_buffer.append(_AppendixItem(text=block.text))
            return

        # lede style -> drop cap paragraph (SS5.4)
        if block.style == "lead":
            E.build_lede(self._doc, block.text, theme=self._theme)
            return

        E.build_narrative(self._doc, block.text, theme=self._theme)

    def emit_table(self, block: TableBlock, asset: Asset) -> None:
        if isinstance(asset, StatsAsset):
            E.build_stats_table(self._doc, asset.summary_stats)
        elif isinstance(asset, TableAsset):
            df = (
                pd.DataFrame.from_records(asset.df_records)
                if asset.df_records
                else pd.DataFrame()
            )
            heading = block.caption or "\u6570\u636E\u660E\u7EC6"
            E.build_dataframe_table(
                self._doc, df,
                highlight_rules=block.highlight_rules,
                theme=self._theme,
                heading=heading,
            )

    def emit_chart(self, block: ChartBlock, asset: Asset) -> None:
        """Render a chart by embedding a matplotlib PNG (SS5.7).

        Uses block.title/subtitle/source from ChartBlock (not chart
        option). Falls back to data table on render failure.
        Includes OOM retry for matplotlib.
        """
        option = getattr(asset, "option", None)
        if not isinstance(option, dict):
            return

        # SS5.7: Use block fields for title/subtitle/source
        chart_title = block.title or ""
        chart_subtitle = block.subtitle or ""
        chart_source = block.source or ""

        # Try native picture with OOM retry
        png_bytes = self._render_chart_with_retry(option)

        if png_bytes:
            from docx.shared import Inches

            # Figure title (H3) / subtitle (Source style)
            if chart_title:
                self._doc.add_heading(chart_title, level=3)
            if chart_subtitle:
                p_sub = self._doc.add_paragraph(style="Source")
                run_sub = p_sub.add_run(chart_subtitle)
                run_sub.font.size = Pt(self._theme.size_small)
                run_sub.font.italic = False
                run_sub.font.color.rgb = _rgb(self._theme.secondary)

            self._doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))

            if chart_source:
                p_src = self._doc.add_paragraph(style="Source")
                run_src = p_src.add_run(chart_source)
                run_src.font.size = Pt(self._theme.size_small)
                run_src.font.italic = True
                run_src.font.color.rgb = _rgb(self._theme.neutral)
            else:
                self._doc.add_paragraph("")
            return

        # Fallback: data table
        E.build_chart_data_table(self._doc, option)

    def _render_chart_with_retry(self, option: dict) -> bytes | None:
        """Render chart to PNG with OOM retry (SS12.2)."""
        try:
            return self._render_chart_png(option)
        except Exception as e:
            logger.warning(
                "matplotlib render failed (%s); falling back to data table", e,
            )
            return None

    @_retry_renderer_chart
    def _render_chart_png(self, option: dict) -> bytes:
        """Single chart render; retried by decorator on MemoryError."""
        from backend.tools.report._chart_renderer import render_chart_to_png
        return render_chart_to_png(option, self._theme)

    def emit_chart_table_pair(
        self,
        block: ChartTablePairBlock,
        chart_asset: Asset,
        table_asset: Asset,
    ) -> None:
        """Render chart + optional KPI strip; push table to section appendix.

        Per SS5.5: table data is deferred to the section-end
        '完整数据' sub-section (editorial data journal layout).
        """
        # Use block's title/subtitle/source for the chart
        synth_chart = ChartBlock(
            block_id=block.block_id,
            asset_id=block.chart_asset_id,
            title=block.title,
            subtitle=block.subtitle,
            source=block.source,
        )
        self.emit_chart(synth_chart, chart_asset)

        # Emit KPI strip if present (SS5.3)
        if block.kpi_strip and block.kpi_strip.items:
            self.emit_kpi_strip(block.kpi_strip)

        # Push table to section appendix buffer (SS5.5)
        if isinstance(table_asset, TableAsset):
            self._buffer_table_for_appendix(block, table_asset)

    def _buffer_table_for_appendix(
        self,
        block: ChartTablePairBlock,
        table_asset: TableAsset,
    ) -> None:
        """Push table to the current section's appendix buffer."""
        key = self._current_section_name or "_default"
        if key not in self._section_appendix_buffer:
            self._section_appendix_buffer[key] = []
        self._section_appendix_buffer[key].append(
            _AppendixTableEntry(block=block, asset=table_asset)
        )

    def emit_comparison_grid(self, block: ComparisonGridBlock) -> None:
        """Phase 3.2 — N-column comparison grid with primary-coloured
        title row and bulleted item cells."""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Pt

        if not block.columns:
            return
        n_cols = len(block.columns)
        max_items = max(len(c.items) for c in block.columns)
        if max_items == 0:
            return

        table = self._doc.add_table(rows=1 + max_items, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for ci, col in enumerate(block.columns):
            cell = table.rows[0].cells[ci]
            cell.text = col.title
            E._set_cell_shading(cell, self._theme.hex_primary)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(self._theme.size_table_header)
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(*self._theme.white)

        for ci, col in enumerate(block.columns):
            for ri, item in enumerate(col.items):
                cell = table.rows[1 + ri].cells[ci]
                cell.text = f"\u2022 {item}"
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(self._theme.size_body)
        self._doc.add_paragraph("")

    def emit_growth_indicators(self, block: GrowthIndicatorsBlock) -> None:
        if block.growth_rates:
            E.build_growth_indicators(self._doc, block.growth_rates)

    def emit_section_cover(self, block: SectionCoverBlock) -> None:
        """Render section heading in Liangang Data Journal style (SS5.6)."""
        E.build_section_heading(
            self._doc, block.index, block.title, theme=self._theme,
        )


def _rgb(t: tuple[int, int, int]):
    from docx.shared import RGBColor
    return RGBColor(*t)


def Pt(val):
    from docx.shared import Pt as _Pt
    return _Pt(val)
