"""DOCX Report Generation Skill — Skill mode (LLM + tools) with deterministic fallback.

Orchestrator: delegates metadata/content extraction to ``_content_collector``,
then either runs an LLM agent loop that composes the document using tools
(wrapping ``_docx_elements``), or falls back to a deterministic builder.
"""
from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document

from backend.tools.base import BaseSkill, SkillCategory, SkillInput, SkillOutput
from backend.tools.registry import register_skill
from backend.tools.report._content_collector import (
    collect_and_associate,
    NarrativeItem, StatsTableItem, GrowthItem,
    ChartDataItem, DataFrameItem, ReportContent,
)
from backend.tools.report import _docx_elements as E
from backend.tools.report._kpi_extractor import extract_kpis_llm

logger = logging.getLogger("analytica.tools.report_docx")


# ---------------------------------------------------------------------------
# Deterministic builder (extracted from the original execute body)
# ---------------------------------------------------------------------------

def _build_docx_deterministic(doc: Document, report: ReportContent) -> None:
    """Build DOCX content using hardcoded element ordering — the fallback path."""
    E.build_cover_page(doc, report.title, report.author, report.date)
    E.build_toc_placeholder(doc)

    # Batch 4: render KPI cards between TOC and first section
    if report.kpi_cards:
        E.build_kpi_row(doc, report.kpi_cards)

    for section in report.sections:
        # No auto-numbering — section names already include Chinese prefixes
        E.build_section_heading(doc, 0, section.name)
        for item in section.items:
            if isinstance(item, NarrativeItem):
                E.build_narrative(doc, item.text)
            elif isinstance(item, StatsTableItem):
                E.build_stats_table(doc, item.summary_stats)
            elif isinstance(item, GrowthItem):
                E.build_growth_indicators(doc, item.growth_rates)
            elif isinstance(item, DataFrameItem):
                E.build_dataframe_table(doc, item.df)
            elif isinstance(item, ChartDataItem):
                E.build_chart_data_table(doc, item.option)

    E.build_summary_section(doc, report.summary_items)


# ---------------------------------------------------------------------------
# Skill
# ---------------------------------------------------------------------------

@register_skill("skill_report_docx", SkillCategory.REPORT, "Word 报告生成（.docx 文件）",
                input_spec="report_metadata + report_structure + 上游数据/图表引用",
                output_spec="DOCX 文件路径")
class DocxReportSkill(BaseSkill):

    async def execute(self, inp: SkillInput, context: dict[str, Any]) -> SkillOutput:
        try:
            intent = inp.params.get("intent", "")
            task_id = inp.params.get("__task_id__", "")

            report = collect_and_associate(
                inp.params, context,
                task_order=inp.params.get("_task_order"),
            )
            report.kpi_cards = await extract_kpis_llm(
                intent, context, span_emit=inp.span_emit, task_id=task_id,
            )

            # ── Try Skill mode (LLM agent loop) ──
            mode = "deterministic_fallback"
            doc = Document()
            E.build_styles(doc)
            E.build_page_header_footer(doc, report.title)

            try:
                from backend.config import get_settings

                settings = get_settings()
                if not settings.REPORT_AGENT_ENABLED:
                    raise RuntimeError("agent mode disabled by config")

                from langchain_openai import ChatOpenAI

                from backend.tools.report._agent_loop import run_report_agent, serialize_report_content
                from backend.tools.report._docx_tools import DOCX_SYSTEM_PROMPT, make_docx_tools

                llm = ChatOpenAI(
                    base_url=settings.QWEN_API_BASE,
                    api_key=settings.QWEN_API_KEY,
                    model=settings.QWEN_MODEL,
                    temperature=0.2,
                    request_timeout=90,
                    extra_body={"enable_thinking": False},
                )

                tools = make_docx_tools(doc, report)
                user_message = serialize_report_content(report)
                success = await run_report_agent(llm, tools, DOCX_SYSTEM_PROMPT, user_message)

                if success:
                    mode = "llm_agent"
                else:
                    logger.warning("DOCX agent did not finalise; falling back to deterministic")
                    doc = Document()
                    E.build_styles(doc)
                    E.build_page_header_footer(doc, report.title)
                    _build_docx_deterministic(doc, report)

            except Exception as agent_err:
                logger.warning("DOCX agent loop failed (%s); falling back to deterministic", agent_err)
                doc = Document()
                E.build_styles(doc)
                E.build_page_header_footer(doc, report.title)
                _build_docx_deterministic(doc, report)
                mode = "deterministic_fallback_error"

            # ── Serialise ──
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()

            return SkillOutput(
                skill_id=self.skill_id,
                status="success",
                output_type="file",
                data=docx_bytes,
                metadata={
                    "format": "docx",
                    "title": report.title,
                    "file_size_bytes": len(docx_bytes),
                    "mode": mode,
                },
            )

        except Exception as e:
            logger.exception("DOCX generation failed: %s", e)
            return self._fail(str(e))
